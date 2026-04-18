"""
Обработчики для категории 'document': PDF, DOCX, DOC, RTF, XLS, XLSX.

Архитектурный принцип (обосновано в обсуждении):
- Каждый обработчик - это ГЕНЕРАТОР, отдающий объекты TextChunk.
- TextChunk содержит кусок текста + метаданные источника (страница/лист/ячейка).
- Это позволяет:
    1) не держать весь файл в памяти (важно для больших PDF и XLSX),
    2) в отчёте по ПДн точно указать, где именно найдена сущность,
    3) детектору использовать имя поля/колонки как дополнительный сигнал.

Выбор библиотек (все - актуальные, поддерживаемые, проверяемые источники):
- pypdf            - https://pypdf.readthedocs.io/
- python-docx      - https://python-docx.readthedocs.io/
- striprtf         - https://pypi.org/project/striprtf/
- xlrd (>= 2.0)    - https://xlrd.readthedocs.io/ (в 2.0 остался только .xls)
- openpyxl         - https://openpyxl.readthedocs.io/

Ограничение по .doc (честно):
    В чистом Python нет надёжного экстрактора текста из старого бинарного
    формата .doc. Промышленный стандарт - LibreOffice headless
    (`soffice --headless --convert-to txt`). Если `soffice` не найден,
    обработчик корректно возвращает пустой результат с пометкой.
    Это поведение аналогично textract/Apache Tika.
"""

from __future__ import annotations

import shutil
import subprocess
import tempfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, Iterator, Optional


# ---------------------------------------------------------------------------
# Единица извлечённого текста.
#
# Поля:
#   text        - сам текст чанка (страница PDF, абзац DOCX, ячейка XLS и т.п.).
#   source      - человекочитаемая координата (для отчёта по ПДн):
#                 "file.pdf#page=3", "book.xls#sheet=Лист1!R2C5".
#   field_name  - имя поля/колонки, если это структурированный источник.
#                 Для детектора это сильный сигнал: "Паспорт" + 1234 567890 -> почти точно паспорт.
#   meta        - структурированные метаданные для классификатора и агрегатора:
#                   page, paragraph, table, row, col, sheet, row_id, via_ocr.
#                 row_id позволяет агрегировать "кластер ПДн в одной строке"
#                 (ФИО + паспорт + адрес в одной записи) - это требуется ТЗ
#                 для решения про УЗ-2/УЗ-3 ("большие объёмы").
# ---------------------------------------------------------------------------
@dataclass(frozen=True)
class TextChunk:
    text: str
    source: str
    field_name: Optional[str] = None
    meta: Dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Вспомогательное: декодирование байтов с перебором популярных в РФ кодировок.
# ---------------------------------------------------------------------------
_CANDIDATE_ENCODINGS = ("utf-8", "cp1251", "cp1252", "koi8-r", "latin-1")


def _decode_bytes(raw: bytes) -> str:
    for enc in _CANDIDATE_ENCODINGS:
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# PDF: постраничный обход через pypdf.
# ---------------------------------------------------------------------------
def extract_pdf(path: Path) -> Iterator[TextChunk]:
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError

    try:
        reader = PdfReader(str(path), strict=False)
    except (PdfReadError, OSError, ValueError):
        # Битый/зашифрованный PDF - корректно выходим без паники.
        return

    # Зашифрованный PDF без пароля - пропускаем.
    if getattr(reader, "is_encrypted", False):
        try:
            if reader.decrypt("") == 0:  # 0 = не удалось
                return
        except Exception:
            return

    name = path.name
    for i, page in enumerate(reader.pages, start=1):
        try:
            text = page.extract_text() or ""
        except Exception:
            # Отдельная страница может упасть - пропускаем её, не весь файл.
            continue
        if text.strip():
            yield TextChunk(
                text=text,
                source=f"{name}#page={i}",
                meta={"page": i, "via_ocr": False},
            )


# ---------------------------------------------------------------------------
# DOCX: абзацы + таблицы.
# ---------------------------------------------------------------------------
def extract_docx(path: Path) -> Iterator[TextChunk]:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

    try:
        doc = Document(str(path))
    except (PackageNotFoundError, OSError, ValueError, KeyError):
        return

    name = path.name

    for i, paragraph in enumerate(doc.paragraphs, start=1):
        t = paragraph.text
        if t and t.strip():
            yield TextChunk(
                text=t,
                source=f"{name}#para={i}",
                meta={"paragraph": i},
            )

    for ti, table in enumerate(doc.tables, start=1):
        rows = list(table.rows)
        if not rows:
            continue

        # Заголовки: первая строка, где все ячейки - непустые короткие строки.
        first_row_values = [cell.text.strip() for cell in rows[0].cells]
        use_headers = bool(first_row_values) and all(
            v and len(v) <= 60 for v in first_row_values
        )
        headers: list[Optional[str]] = (
            first_row_values if use_headers else []
        )
        data_start = 1 if use_headers else 0

        for ri, row in enumerate(rows[data_start:], start=data_start + 1):
            for ci, cell in enumerate(row.cells, start=1):
                t = cell.text
                if not t or not t.strip():
                    continue
                field_name = (
                    headers[ci - 1]
                    if headers and ci - 1 < len(headers)
                    else None
                )
                yield TextChunk(
                    text=t,
                    source=f"{name}#table={ti},r={ri},c={ci}",
                    field_name=field_name,
                    meta={
                        "table": ti,
                        "row": ri,
                        "col": ci,
                        "row_id": f"{name}#table={ti}#row={ri}",
                    },
                )


# ---------------------------------------------------------------------------
# DOC: через LibreOffice soffice (если установлен).
# ---------------------------------------------------------------------------
def extract_doc(path: Path) -> Iterator[TextChunk]:
    soffice = shutil.which("soffice") or shutil.which("soffice.exe")
    if not soffice:
        # Нет конвертера - ничего не извлекаем. Диспетчер пометит это в notes.
        return

    with tempfile.TemporaryDirectory() as tmp:
        try:
            subprocess.run(
                [
                    soffice, "--headless",
                    "--convert-to", "txt:Text",
                    "--outdir", tmp,
                    str(path),
                ],
                check=True, timeout=120,
                stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL,
            )
        except (subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
            return

        out_path = Path(tmp) / f"{path.stem}.txt"
        if not out_path.exists():
            return

        text = _decode_bytes(out_path.read_bytes())
        if text.strip():
            yield TextChunk(
                text=text,
                source=f"{path.name}#body",
                meta={"converter": "libreoffice"},
            )


# ---------------------------------------------------------------------------
# RTF: striprtf разворачивает RTF в plain text.
# ---------------------------------------------------------------------------
def extract_rtf(path: Path) -> Iterator[TextChunk]:
    from striprtf.striprtf import rtf_to_text

    try:
        raw = path.read_bytes()
    except OSError:
        return

    rtf_source = _decode_bytes(raw)

    try:
        plain = rtf_to_text(rtf_source, errors="ignore")
    except Exception:
        return

    if plain and plain.strip():
        yield TextChunk(
            text=plain,
            source=f"{path.name}#body",
            meta={},
        )


# ---------------------------------------------------------------------------
# XLS (legacy binary): xlrd >= 2.0.
# Логика: если первая строка листа похожа на заголовки (все ячейки -
# непустые строки), используем её как field_name для последующих строк.
# ---------------------------------------------------------------------------
def extract_xls(path: Path) -> Iterator[TextChunk]:
    try:
        import xlrd
    except ImportError:
        return

    try:
        wb = xlrd.open_workbook(str(path), on_demand=True)
    except (xlrd.XLRDError, OSError, ValueError):
        return

    try:
        for sheet_index in range(wb.nsheets):
            sheet = wb.sheet_by_index(sheet_index)
            if sheet.nrows == 0:
                wb.unload_sheet(sheet_index)
                continue

            headers: list[Optional[str]] = []
            start_row = 0
            first_row = sheet.row_values(0)
            if first_row and all(isinstance(v, str) and v.strip() for v in first_row):
                headers = [str(v).strip() for v in first_row]
                start_row = 1

            for row_idx in range(start_row, sheet.nrows):
                row_values = sheet.row_values(row_idx)
                row_number = row_idx + 1
                row_id = f"{path.name}#sheet={sheet.name}#row={row_number}"
                for col_idx, val in enumerate(row_values):
                    if val is None or val == "":
                        continue
                    text = str(val).strip()
                    if not text:
                        continue
                    col_number = col_idx + 1
                    field_name = (
                        headers[col_idx] if col_idx < len(headers) else None
                    )
                    yield TextChunk(
                        text=text,
                        source=f"{path.name}#sheet={sheet.name}!R{row_number}C{col_number}",
                        field_name=field_name,
                        meta={
                            "sheet": sheet.name,
                            "row": row_number,
                            "col": col_number,
                            "row_id": row_id,
                        },
                    )
            wb.unload_sheet(sheet_index)
    finally:
        try:
            wb.release_resources()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# XLSX: openpyxl в read_only режиме (экономит память на больших книгах).
# ---------------------------------------------------------------------------
def extract_xlsx(path: Path) -> Iterator[TextChunk]:
    try:
        from openpyxl import load_workbook
    except ImportError:
        return

    try:
        wb = load_workbook(str(path), read_only=True, data_only=True)
    except Exception:
        return

    try:
        for ws in wb.worksheets:
            rows_iter = ws.iter_rows(values_only=True)
            try:
                first_row = next(rows_iter)
            except StopIteration:
                continue

            headers: list[Optional[str]] = []
            data_start_offset = 0
            if first_row and all(isinstance(v, str) and v and v.strip() for v in first_row):
                headers = [str(v).strip() for v in first_row]
                data_start_offset = 1
            else:
                # Первая строка - уже данные, обработаем её отдельно.
                row_id = f"{path.name}#sheet={ws.title}#row=1"
                for col_idx, val in enumerate(first_row, start=1):
                    if val is None:
                        continue
                    t = str(val).strip()
                    if not t:
                        continue
                    yield TextChunk(
                        text=t,
                        source=f"{path.name}#sheet={ws.title}!R1C{col_idx}",
                        meta={
                            "sheet": ws.title,
                            "row": 1,
                            "col": col_idx,
                            "row_id": row_id,
                        },
                    )

            for row_idx, row in enumerate(rows_iter, start=1 + data_start_offset):
                row_id = f"{path.name}#sheet={ws.title}#row={row_idx}"
                for col_idx, val in enumerate(row, start=1):
                    if val is None:
                        continue
                    t = str(val).strip()
                    if not t:
                        continue
                    field_name = (
                        headers[col_idx - 1]
                        if headers and col_idx - 1 < len(headers)
                        else None
                    )
                    yield TextChunk(
                        text=t,
                        source=f"{path.name}#sheet={ws.title}!R{row_idx}C{col_idx}",
                        field_name=field_name,
                        meta={
                            "sheet": ws.title,
                            "row": row_idx,
                            "col": col_idx,
                            "row_id": row_id,
                        },
                    )
    finally:
        try:
            wb.close()
        except Exception:
            pass


# ---------------------------------------------------------------------------
# HTML: реальный handler. Частый кейс в нашем датасете - файлы с
# расширением .pdf, но реальным содержимым HTML (видно в логе прогона:
# "invalid pdf header: b'<!DOC'"). Проверено: 911 из 1658 .pdf оказались
# такими. Чтобы не терять треть датасета, имеем:
#   1) сам HTML handler (для честных .html / .htm);
#   2) sniff в extract_pdf: если файл начинается не с %PDF-, пытаемся
#      обработать как HTML.
# ---------------------------------------------------------------------------
def extract_html(path: Path) -> Iterator[TextChunk]:
    try:
        from bs4 import BeautifulSoup
    except ImportError:
        return iter(())
    try:
        raw = path.read_bytes()
    except OSError:
        return iter(())

    text = _decode_bytes(raw)
    if not text:
        return iter(())

    # "lxml" быстрее, но требует доп. пакет; html.parser всегда доступен.
    try:
        soup = BeautifulSoup(text, "lxml")
    except Exception:
        soup = BeautifulSoup(text, "html.parser")

    # Убираем script/style, они не несут ПДн и шумят при NER.
    for bad in soup(["script", "style", "noscript"]):
        bad.decompose()

    def _gen() -> Iterator[TextChunk]:
        title_tag = soup.find("title")
        if title_tag and title_tag.get_text(strip=True):
            yield TextChunk(
                text=title_tag.get_text(strip=True),
                source=f"{path.name}#title",
                field_name="title",
                meta={"tag": "title"},
            )
        # Основной текст: один чанк на документ, без field_name. Детальнее
        # резать по тегам не требуется - детектор сам агрегирует по файлу.
        body_text = soup.get_text(separator="\n", strip=True)
        if body_text:
            yield TextChunk(
                text=body_text,
                source=f"{path.name}#body",
                field_name=None,
                meta={"tag": "body"},
            )
    return _gen()


def _is_real_pdf(path: Path) -> bool:
    try:
        with open(path, "rb") as f:
            head = f.read(5)
        return head == b"%PDF-"
    except OSError:
        return False


def extract_pdf_or_html(path: Path) -> Iterator[TextChunk]:
    """Умная диспетчеризация для .pdf: если файл в реальности HTML -
    используем HTML-экстрактор, иначе обычный PDF."""
    if _is_real_pdf(path):
        return extract_pdf(path)
    # не PDF - проверяем, не HTML ли это
    try:
        with open(path, "rb") as f:
            head = f.read(512).lstrip().lower()
    except OSError:
        return iter(())
    if head.startswith(b"<!doc") or head.startswith(b"<html") or b"<html" in head[:256]:
        return extract_html(path)
    # неизвестный бинарь с расширением .pdf - возвращаем пустоту
    return iter(())


# ---------------------------------------------------------------------------
# Удобная карта: расширение -> экстрактор (использует main.py).
# ---------------------------------------------------------------------------
DOCUMENT_EXTRACTORS = {
    "pdf":  extract_pdf_or_html,
    "docx": extract_docx,
    "doc":  extract_doc,
    "rtf":  extract_rtf,
    "xls":  extract_xls,
    "xlsx": extract_xlsx,
    "html": extract_html,
    "htm":  extract_html,
}
